"""
Export Service ‚Äî Xu·∫•t ƒë·ªÅ thi d·∫°ng DOCX, LaTeX, PDF (HTML chu·∫©n in).

H·ªó tr·ª£:
  - Xu·∫•t c√¢u h·ªèi sinh b·ªüi AI (list dict)
  - Xu·∫•t c√¢u h·ªèi t·ª´ ng√¢n h√†ng (list DB objects)
  - Gom theo m·ª©c ƒë·ªô ho·∫∑c danh s√°ch ph·∫≥ng
  - C√≥/kh√¥ng ƒë√°p √°n + l·ªùi gi·∫£i

Quy ∆∞·ªõc LaTeX trong question_text / answer / solution_steps:
  - Inline math: $...$ ho·∫∑c \\(...\\)
  - Display math: $$...$$ ho·∫∑c \\[...\\]
"""

import io
import json
import re
import os
import tempfile
from typing import List, Dict, Optional
from datetime import datetime

from docx import Document as DocxDocument
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.services.latex_to_omml import add_math_to_paragraph

# ‚îÄ‚îÄ‚îÄ Difficulty helpers ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ‚îÄ
DIFF_LABELS = {
    "NB": "Nh·∫≠n bi·∫øt",
    "TH": "Th√¥ng hi·ªÉu",
    "VD": "V·∫≠n d·ª•ng",
    "VDC": "V·∫≠n d·ª•ng cao",
}
DIFF_ORDER = ["NB", "TH", "VD", "VDC"]


def _normalize_questions(questions) -> List[Dict]:
    """Convert DB objects or raw dicts to uniform list of dicts."""
    out = []
    for q in questions:
        if isinstance(q, dict):
            d = q.copy()
            # Normalize key names (generated uses 'question', DB uses 'question_text')
            if "question_text" in d and "question" not in d:
                d["question"] = d.pop("question_text")
            if "question_type" in d and "type" not in d:
                d["type"] = d.pop("question_type")
            # solution_steps may be JSON string from DB
            steps = d.get("solution_steps", [])
            if isinstance(steps, str):
                try:
                    steps = json.loads(steps)
                except Exception:
                    steps = [steps] if steps.strip() else []
            d["solution_steps"] = steps if isinstance(steps, list) else []
            out.append(d)
        else:
            # SQLAlchemy model
            steps = q.solution_steps or "[]"
            if isinstance(steps, str):
                try:
                    steps = json.loads(steps)
                except Exception:
                    steps = [steps] if steps.strip() else []
            out.append({
                "question": q.question_text,
                "type": q.question_type or "",
                "topic": q.topic or "",
                "difficulty": q.difficulty or "TH",
                "grade": q.grade,
                "chapter": q.chapter or "",
                "lesson_title": q.lesson_title or "",
                "answer": q.answer or "",
                "solution_steps": steps if isinstance(steps, list) else [],
            })

    # Sanitize all text fields ‚Äî remove XML-invalid control chars
    for d in out:
        for key in ("question", "type", "topic", "difficulty", "answer"):
            if key in d and isinstance(d[key], str):
                d[key] = _sanitize_for_xml(d[key])
        if "solution_steps" in d and isinstance(d["solution_steps"], list):
            d["solution_steps"] = [
                _sanitize_for_xml(s) if isinstance(s, str) else s
                for s in d["solution_steps"]
            ]

    return out


def _group_by_difficulty(questions: List[Dict]) -> Dict[str, List[Dict]]:
    """Group questions by difficulty in standard order."""
    groups = {}
    for q in questions:
        d = q.get("difficulty", "TH")
        groups.setdefault(d, []).append(q)
    return groups


def _strip_latex_delimiters(text: str) -> str:
    """Clean LaTeX for plain-text contexts (DOCX)."""
    if not text:
        return ""
    # Keep the math content but remove outer delimiters for inline display
    return text


def _sanitize_for_xml(text: str) -> str:
    """Remove characters that are invalid in XML (python-docx requirement).
    
    XML 1.0 allows: #x9 | #xA | #xD | [#x20-#xD7FF] | [#xE000-#xFFFD]
    Everything else (NULL, control chars 0x01-0x08, 0x0B-0x0C, 0x0E-0x1F) must go.
    """
    if not text:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)


# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
#  DOCX EXPORT
# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê

def _set_cell_shading(cell, color: str):
    """Set table cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn("w:shd"), {
        qn("w:fill"): color,
        qn("w:val"): "clear",
    })
    shading.append(shading_elm)


def export_docx(
    questions,
    title: str = "ƒê·ªÄ THI TO√ÅN H·ªåC",
    subtitle: str = "",
    include_answers: bool = True,
    include_solutions: bool = True,
    group_by_diff: bool = True,
    exam_info: Optional[Dict] = None,
) -> io.BytesIO:
    """
    Generate a professional DOCX exam document.

    Args:
        questions: list of dicts or DB objects
        title: exam title
        subtitle: topic or subtitle
        include_answers: include answers
        include_solutions: include solution steps
        group_by_diff: group questions by difficulty level
        exam_info: optional dict with extra info (date, time_limit, etc.)

    Returns: BytesIO buffer containing the DOCX file
    """
    items = _normalize_questions(questions)
    doc = DocxDocument()

    # Sanitize text inputs
    title = _sanitize_for_xml(title)
    subtitle = _sanitize_for_xml(subtitle)

    # ‚îÄ‚îÄ Page setup ‚îÄ‚îÄ
    section = doc.sections[0]
    section.page_width = Cm(21)     # A4
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    # ‚îÄ‚îÄ Styles ‚îÄ‚îÄ
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3

    # ‚îÄ‚îÄ Header block ‚îÄ‚îÄ
    # School / Organization line
    info = exam_info or {}
    org = info.get("organization", "TR∆Ø·ªúNG THPT ................")
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_org.add_run(org)
    run.font.size = Pt(12)
    run.font.bold = True

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(8)
    run = p_title.add_run(title)
    run.font.size = Pt(16)
    run.font.bold = True

    # Subtitle (topic)
    if subtitle:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_sub.add_run(subtitle)
        run.font.size = Pt(13)
        run.font.bold = True

    # Info table (Date, Time, etc.)
    date_str = info.get("date", datetime.now().strftime("%d/%m/%Y"))
    time_limit = info.get("time_limit", "")
    total_q = len(items)

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_before = Pt(4)
    p_info.paragraph_format.space_after = Pt(4)
    info_parts = [f"Ng√†y: {date_str}", f"S·ªë c√¢u: {total_q}"]
    if time_limit:
        info_parts.append(f"Th·ªùi gian: {time_limit}")
    run = p_info.add_run("  |  ".join(info_parts))
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Separator line
    p_sep = doc.add_paragraph()
    p_sep.paragraph_format.space_before = Pt(2)
    p_sep.paragraph_format.space_after = Pt(12)
    run = p_sep.add_run("‚îÄ" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(180, 180, 180)
    p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ‚îÄ‚îÄ Questions ‚îÄ‚îÄ
    num = 0

    def _write_question(q: Dict, num: int):
        """Write a single question to the document with OMML math rendering."""
        diff = q.get("difficulty", "")
        diff_label = DIFF_LABELS.get(diff, diff)

        # Question header: "C√¢u 1. [TH]"
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_before = Pt(10)
        p_q.paragraph_format.keep_with_next = True

        run = p_q.add_run(f"C√¢u {num}.")
        run.font.bold = True
        run.font.size = Pt(12)

        if diff_label:
            run = p_q.add_run(f"  [{diff_label}]")
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)

        # Question body ‚Äî render LaTeX math as OMML equations
        q_text = q.get("question", "")
        for line in q_text.split("\n"):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(2)
            add_math_to_paragraph(p, line, font_size=24)  # 24 half-pts = 12pt

        # Answer ‚Äî render math in answer text too
        if include_answers and q.get("answer"):
            p_ans = doc.add_paragraph()
            p_ans.paragraph_format.left_indent = Cm(0.5)
            p_ans.paragraph_format.space_before = Pt(6)

            run = p_ans.add_run("ƒê√°p √°n: ")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0, 128, 80)

            add_math_to_paragraph(p_ans, q["answer"], font_size=22, font_color="008050")

        # Solution steps ‚Äî render math in each step
        if include_solutions and q.get("solution_steps"):
            p_sol_header = doc.add_paragraph()
            p_sol_header.paragraph_format.left_indent = Cm(0.5)
            p_sol_header.paragraph_format.space_before = Pt(4)
            run = p_sol_header.add_run("L·ªùi gi·∫£i:")
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(180, 120, 0)

            for i, step in enumerate(q["solution_steps"], 1):
                p_step = doc.add_paragraph()
                p_step.paragraph_format.left_indent = Cm(1)
                p_step.paragraph_format.space_after = Pt(2)

                run = p_step.add_run(f"B∆∞·ªõc {i}: ")
                run.font.bold = True
                run.font.size = Pt(11)

                add_math_to_paragraph(p_step, step, font_size=22)

    if group_by_diff and len(items) > 1:
        groups = _group_by_difficulty(items)
        for diff_key in DIFF_ORDER:
            if diff_key not in groups:
                continue
            group = groups[diff_key]
            label = DIFF_LABELS.get(diff_key, diff_key)

            # Section header
            p_sec = doc.add_paragraph()
            p_sec.paragraph_format.space_before = Pt(16)
            p_sec.paragraph_format.space_after = Pt(6)
            p_sec.paragraph_format.keep_with_next = True

            # Create a simple shaded paragraph for section header
            run = p_sec.add_run(f"‚ñå {label} ({len(group)} c√¢u)")
            run.font.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(50, 50, 80)

            for q in group:
                num += 1
                _write_question(q, num)
    else:
        for q in items:
            num += 1
            _write_question(q, num)

    # ‚îÄ‚îÄ Footer ‚îÄ‚îÄ
    p_footer = doc.add_paragraph()
    p_footer.paragraph_format.space_before = Pt(24)
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_footer.add_run("‚îÄ‚îÄ H·∫æT ‚îÄ‚îÄ")
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(150, 150, 150)

    p_gen = doc.add_paragraph()
    p_gen.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_gen.add_run("T·∫°o b·ªüi Math Exam Parser AI")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = RGBColor(180, 180, 180)

    # Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
#  DOCX ‚Äî CH·ªà ƒê·ªÄ (kh√¥ng ƒë√°p √°n) + ƒê√ÅP √ÅN RI√äNG
# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê

def export_docx_split(
    questions,
    title: str = "ƒê·ªÄ THI TO√ÅN H·ªåC",
    subtitle: str = "",
    exam_info: Optional[Dict] = None,
) -> Dict[str, io.BytesIO]:
    """
    Export DOCX t√°ch ƒë·ªÅ v√† ƒë√°p √°n ri√™ng bi·ªát.

    Returns dict:
        "exam": BytesIO  ‚Äî file ƒë·ªÅ (kh√¥ng ƒë√°p √°n)
        "answers": BytesIO ‚Äî file ƒë√°p √°n + l·ªùi gi·∫£i
    """
    exam_buf = export_docx(
        questions, title=title, subtitle=subtitle,
        include_answers=False, include_solutions=False,
        exam_info=exam_info,
    )
    answer_buf = export_docx(
        questions, title=f"ƒê√ÅP √ÅN - {title}", subtitle=subtitle,
        include_answers=True, include_solutions=True,
        exam_info=exam_info,
    )
    return {"exam": exam_buf, "answers": answer_buf}


# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
#  LATEX EXPORT
# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê

def _escape_latex(text: str) -> str:
    """Escape special LaTeX characters OUTSIDE of math delimiters."""
    if not text:
        return ""

    # Protect math regions
    parts = []
    last = 0
    # Match $...$ and $$...$$ and \(...\) and \[...\]
    pattern = re.compile(r'(\$\$.*?\$\$|\$.*?\$|\\\(.*?\\\)|\\\[.*?\\\])', re.DOTALL)
    for m in pattern.finditer(text):
        # Escape the non-math part
        before = text[last:m.start()]
        before = _escape_latex_chars(before)
        parts.append(before)
        parts.append(m.group())  # Keep math as-is
        last = m.end()
    # Remaining text
    parts.append(_escape_latex_chars(text[last:]))
    return "".join(parts)


def _escape_latex_chars(text: str) -> str:
    """Escape LaTeX special chars in non-math text.
    
    Order matters: escape backslash LAST to avoid double-escaping.
    """
    if not text:
        return ""
    # First pass: escape all special chars EXCEPT backslash
    for old, new in [
        ("&", "\\&"),
        ("%", "\\%"),
        ("#", "\\#"),
        ("_", "\\_"),
        ("~", "\\textasciitilde{}"),
        ("^", "\\textasciicircum{}"),
    ]:
        text = text.replace(old, new)
    # Don't escape { } and \ ‚Äî they break LaTeX structure
    # If there are literal braces/backslashes outside math, leave as-is
    return text


def export_latex(
    questions,
    title: str = "ƒê·ªÄ THI TO√ÅN H·ªåC",
    subtitle: str = "",
    include_answers: bool = True,
    include_solutions: bool = True,
    group_by_diff: bool = True,
    exam_info: Optional[Dict] = None,
) -> io.BytesIO:
    """
    Generate a LaTeX .tex file for the exam.

    Returns: BytesIO buffer containing UTF-8 .tex content
    """
    items = _normalize_questions(questions)
    info = exam_info or {}
    date_str = info.get("date", datetime.now().strftime("%d/%m/%Y"))
    time_limit = info.get("time_limit", "")
    org = info.get("organization", "TR∆Ø·ªúNG THPT ................")
    total = len(items)

    lines = []
    lines.append(r"\documentclass[12pt,a4paper]{article}")
    lines.append("")
    lines.append(r"% ‚îÄ‚îÄ Packages ‚îÄ‚îÄ")
    lines.append(r"\usepackage[utf8]{inputenc}")
    lines.append(r"\usepackage[vietnamese]{babel}")
    lines.append(r"\usepackage{amsmath,amssymb,amsfonts}")
    lines.append(r"\usepackage{geometry}")
    lines.append(r"\usepackage{enumitem}")
    lines.append(r"\usepackage{fancyhdr}")
    lines.append(r"\usepackage{xcolor}")
    lines.append(r"\usepackage{tcolorbox}")
    lines.append(r"\usepackage{tikz}")
    lines.append("")
    lines.append(r"\geometry{top=2cm, bottom=2cm, left=2.5cm, right=2cm}")
    lines.append(r"\pagestyle{fancy}")
    lines.append(r"\fancyhf{}")
    lines.append(r"\fancyfoot[C]{\thepage}")
    lines.append(r"\renewcommand{\headrulewidth}{0pt}")
    lines.append("")
    lines.append(r"% ‚îÄ‚îÄ Custom commands ‚îÄ‚îÄ")
    lines.append(r"\definecolor{answergreen}{RGB}{0,128,80}")
    lines.append(r"\definecolor{solutiongold}{RGB}{180,120,0}")
    lines.append(r"\definecolor{sectionblue}{RGB}{50,50,120}")
    lines.append(r"\definecolor{lightgray}{RGB}{240,240,240}")
    lines.append("")
    lines.append(r"\newcommand{\dapan}[1]{\textcolor{answergreen}{\textbf{ƒê√°p √°n:} #1}}")
    lines.append(r"\newcommand{\loigiai}{\textcolor{solutiongold}{\textbf{L·ªùi gi·∫£i:}}}")
    lines.append(r"\newcommand{\mucdo}[1]{\hfill\textit{\small\textcolor{gray}{[#1]}}}")
    lines.append("")
    lines.append(r"\tcbset{")
    lines.append(r"  sectionbox/.style={colback=lightgray,colframe=sectionblue,")
    lines.append(r"    fonttitle=\bfseries,boxrule=0.5pt,left=8pt,right=8pt,top=4pt,bottom=4pt}")
    lines.append(r"}")
    lines.append("")
    lines.append(r"\begin{document}")
    lines.append("")

    # ‚îÄ‚îÄ Header ‚îÄ‚îÄ
    lines.append(r"\begin{center}")
    lines.append(r"  \textbf{" + _escape_latex(org) + r"} \\[6pt]")
    lines.append(r"  {\LARGE \textbf{" + _escape_latex(title) + r"}} \\[4pt]")
    if subtitle:
        lines.append(r"  {\large \textbf{" + _escape_latex(subtitle) + r"}} \\[4pt]")
    info_parts = [f"Ng√†y: {date_str}", f"S·ªë c√¢u: {total}"]
    if time_limit:
        info_parts.append(f"Th·ªùi gian: {time_limit}")
    lines.append(r"  {\small\textit{" + " $|$ ".join(info_parts) + r"}} \\[2pt]")
    lines.append(r"  \rule{0.8\textwidth}{0.4pt}")
    lines.append(r"\end{center}")
    lines.append(r"\vspace{8pt}")
    lines.append("")

    # ‚îÄ‚îÄ Questions ‚îÄ‚îÄ
    num = 0

    def _write_q_latex(q: Dict, num: int):
        diff = q.get("difficulty", "")
        diff_label = DIFF_LABELS.get(diff, diff)
        q_text = q.get("question", "")

        lines.append("")
        lines.append(r"\noindent\textbf{C√¢u " + str(num) + r".}")
        if diff_label:
            lines.append(r"\mucdo{" + diff_label + r"}")
        lines.append(r"\\")

        # Question body ‚Äî each line
        for line in q_text.split("\n"):
            line = line.strip()
            if line:
                lines.append(r"\indent " + line + r" \\")

        # Answer
        if include_answers and q.get("answer"):
            lines.append(r"\dapan{" + q["answer"] + r"}")
            lines.append(r"\\")

        # Solution
        if include_solutions and q.get("solution_steps"):
            lines.append(r"\loigiai")
            lines.append(r"\begin{enumerate}[leftmargin=2cm, label=\textbf{B∆∞·ªõc \arabic*:}]")
            for step in q["solution_steps"]:
                lines.append(r"  \item " + step)
            lines.append(r"\end{enumerate}")

        lines.append(r"\vspace{6pt}")

    if group_by_diff and len(items) > 1:
        groups = _group_by_difficulty(items)
        for diff_key in DIFF_ORDER:
            if diff_key not in groups:
                continue
            group = groups[diff_key]
            label = DIFF_LABELS.get(diff_key, diff_key)

            lines.append("")
            lines.append(r"\begin{tcolorbox}[sectionbox, title={" + label + f" ({len(group)} c√¢u)" + r"}]")
            lines.append(r"\end{tcolorbox}")
            lines.append(r"\vspace{4pt}")

            for q in group:
                num += 1
                _write_q_latex(q, num)
    else:
        for q in items:
            num += 1
            _write_q_latex(q, num)

    # ‚îÄ‚îÄ Footer ‚îÄ‚îÄ
    lines.append("")
    lines.append(r"\vspace{20pt}")
    lines.append(r"\begin{center}")
    lines.append(r"  \textbf{‚îÄ‚îÄ H·∫æT ‚îÄ‚îÄ}")
    lines.append(r"\end{center}")
    lines.append("")
    lines.append(r"\end{document}")

    content = "\n".join(lines)
    buf = io.BytesIO(content.encode("utf-8"))
    buf.seek(0)
    return buf


# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê
#  PDF EXPORT (Enhanced HTML for print)
# ‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê‚ïê

def export_pdf_html(
    questions,
    title: str = "ƒê·ªÄ THI TO√ÅN H·ªåC",
    subtitle: str = "",
    include_answers: bool = True,
    include_solutions: bool = True,
    group_by_diff: bool = True,
    exam_info: Optional[Dict] = None,
) -> str:
    """
    Generate enhanced HTML optimized for PDF printing.

    Returns: HTML string (rendered client-side with KaTeX, printed to PDF)
    """
    items = _normalize_questions(questions)
    info = exam_info or {}
    date_str = info.get("date", datetime.now().strftime("%d/%m/%Y"))
    time_limit = info.get("time_limit", "")
    org = info.get("organization", "TR∆Ø·ªúNG THPT ................")
    total = len(items)

    def _esc(t):
        if not t:
            return ""
        return (t.replace("&", "&amp;").replace("<", "&lt;")
                 .replace(">", "&gt;").replace("\n", "<br>"))

    questions_html = ""
    num = 0

    def _build_q(q, num):
        diff = q.get("difficulty", "")
        diff_label = DIFF_LABELS.get(diff, diff)
        q_text = q.get("question", "")
        steps = q.get("solution_steps", [])

        h = f'<div class="question">'
        h += f'<div class="q-header"><span class="q-num">C√¢u {num}.</span>'
        if diff_label:
            h += f' <span class="badge badge-{diff.lower()}">{diff_label}</span>'
        h += '</div>'
        h += f'<div class="q-body">{_esc(q_text)}</div>'

        if include_answers and q.get("answer"):
            h += f'<div class="answer-box"><span class="label">ƒê√°p √°n:</span> {_esc(q["answer"])}</div>'

        if include_solutions and steps:
            h += '<div class="solution-box"><span class="label">L·ªùi gi·∫£i:</span>'
            for i, step in enumerate(steps, 1):
                h += f'<div class="step"><b>B∆∞·ªõc {i}:</b> {_esc(step)}</div>'
            h += '</div>'

        h += '</div>'
        return h

    if group_by_diff and len(items) > 1:
        groups = _group_by_difficulty(items)
        for diff_key in DIFF_ORDER:
            if diff_key not in groups:
                continue
            group = groups[diff_key]
            label = DIFF_LABELS.get(diff_key, diff_key)
            questions_html += f'<div class="section-header"><span class="section-marker"></span>{label} ({len(group)} c√¢u)</div>'
            for q in group:
                num += 1
                questions_html += _build_q(q, num)
    else:
        for q in items:
            num += 1
            questions_html += _build_q(q, num)

    info_line = f"Ng√†y: {_esc(date_str)} &nbsp;|&nbsp; S·ªë c√¢u: {total}"
    if time_limit:
        info_line += f" &nbsp;|&nbsp; Th·ªùi gian: {_esc(time_limit)}"

    html = f"""<!DOCTYPE html>
<html lang="vi">
<head>
<meta charset="UTF-8">
<title>{_esc(title)}</title>
<link rel="stylesheet" href="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.css">
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/katex.min.js"></script>
<script src="https://cdn.jsdelivr.net/npm/katex@0.16.9/dist/contrib/auto-render.min.js"></script>
<style>
@page {{
    size: A4;
    margin: 18mm 22mm 18mm 25mm;
}}
* {{ margin: 0; padding: 0; box-sizing: border-box; }}
body {{
    font-family: "Times New Roman", "Noto Serif", serif;
    font-size: 12pt;
    line-height: 1.65;
    color: #1a1a1a;
    padding: 0;
}}
.print-toolbar {{
    position: fixed; top: 0; left: 0; right: 0; z-index: 1000;
    background: linear-gradient(135deg, #4f46e5, #6366f1);
    color: #fff; padding: 12px 24px;
    display: flex; align-items: center; justify-content: center; gap: 16px;
    font-family: -apple-system, sans-serif; font-size: 14px;
    box-shadow: 0 4px 16px rgba(0,0,0,0.15);
}}
.print-toolbar button {{
    background: #fff; color: #4f46e5; border: none; padding: 8px 28px;
    border-radius: 8px; font-weight: 700; cursor: pointer; font-size: 14px;
    transition: transform 0.15s;
}}
.print-toolbar button:hover {{ transform: scale(1.03); }}
.exam-container {{ max-width: 210mm; margin: 60px auto 40px; padding: 0 32px; }}

/* Header */
.exam-header {{
    text-align: center; padding-bottom: 16px;
    border-bottom: 2.5px solid #1a1a1a; margin-bottom: 20px;
}}
.header-org {{ font-size: 12pt; font-weight: 700; text-transform: uppercase; letter-spacing: 0.5px; }}
.header-title {{ font-size: 17pt; font-weight: 700; margin-top: 10px; text-transform: uppercase; letter-spacing: 1px; }}
.header-subtitle {{ font-size: 14pt; font-weight: 700; margin-top: 4px; }}
.header-info {{ font-size: 10pt; color: #555; margin-top: 8px; font-style: italic; }}

/* Sections */
.section-header {{
    margin-top: 22px; padding: 8px 16px;
    background: #f0f0f0; border-left: 4px solid #333;
    font-weight: 700; font-size: 12pt;
    page-break-after: avoid;
    display: flex; align-items: center; gap: 8px;
}}
.section-marker {{
    display: inline-block; width: 6px; height: 6px;
    background: #333; border-radius: 50%;
}}

/* Questions */
.question {{
    margin-bottom: 18px;
    page-break-inside: avoid;
}}
.q-header {{
    margin-bottom: 3px;
    display: flex; align-items: baseline; gap: 8px;
}}
.q-num {{ font-weight: 700; font-size: 12pt; }}
.badge {{
    font-size: 8.5pt; padding: 1px 8px; border-radius: 3px;
    font-weight: 600; letter-spacing: 0.3px;
}}
.badge-nb {{ background: #dbeafe; color: #1e40af; }}
.badge-th {{ background: #d1fae5; color: #065f46; }}
.badge-vd {{ background: #fef3c7; color: #92400e; }}
.badge-vdc {{ background: #fce7f3; color: #9d174d; }}
.q-body {{
    font-size: 12pt; line-height: 1.85;
    margin-left: 20px; margin-top: 2px;
}}

/* Answer */
.answer-box {{
    margin: 6px 0 4px 20px; padding: 8px 14px;
    background: #f0fff4; border-left: 3px solid #22a06b;
    font-size: 11pt;
}}
.answer-box .label {{ font-weight: 700; color: #166534; }}

/* Solution */
.solution-box {{
    margin: 4px 0 0 20px; padding: 8px 14px;
    background: #fffdf0; border-left: 3px solid #d9930e;
    font-size: 11pt;
}}
.solution-box .label {{ font-weight: 700; color: #92400e; }}
.step {{ padding: 2px 0; line-height: 1.7; }}

/* Footer */
.exam-footer {{
    margin-top: 32px; text-align: center;
    border-top: 1.5px solid #ccc; padding-top: 10px;
}}
.exam-footer .end-mark {{ font-weight: 700; font-size: 12pt; color: #555; }}
.exam-footer .generator {{ font-size: 8pt; color: #aaa; font-style: italic; margin-top: 4px; }}

@media print {{
    .print-toolbar {{ display: none !important; }}
    .exam-container {{ margin: 0; padding: 0; max-width: none; }}
    .answer-box {{ background: #f0fff4 !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .solution-box {{ background: #fffdf0 !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .section-header {{ background: #f0f0f0 !important; -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
    .badge {{ -webkit-print-color-adjust: exact; print-color-adjust: exact; }}
}}
</style>
</head>
<body>
<div class="print-toolbar">
    üìÑ Ch·ªçn <strong>"Save as PDF"</strong> ho·∫∑c <strong>"Microsoft Print to PDF"</strong>
    <button onclick="window.print()">In / L∆∞u PDF</button>
</div>
<div class="exam-container">
    <div class="exam-header">
        <div class="header-org">{_esc(org)}</div>
        <div class="header-title">{_esc(title)}</div>
        {"<div class='header-subtitle'>" + _esc(subtitle) + "</div>" if subtitle else ""}
        <div class="header-info">{info_line}</div>
    </div>
    {questions_html}
    <div class="exam-footer">
        <div class="end-mark">‚îÄ‚îÄ H·∫æT ‚îÄ‚îÄ</div>
        <div class="generator">T·∫°o b·ªüi Math Exam Parser AI</div>
    </div>
</div>
<script>
document.addEventListener("DOMContentLoaded", function() {{
    renderMathInElement(document.body, {{
        delimiters: [
            {{left: "$$", right: "$$", display: true}},
            {{left: "$", right: "$", display: false}},
            {{left: "\\\\(", right: "\\\\)", display: false}},
            {{left: "\\\\[", right: "\\\\]", display: true}}
        ],
        throwOnError: false
    }});
}});
</script>
</body>
</html>"""
    return html
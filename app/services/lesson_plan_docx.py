"""
Render KHBD (JSON) → file Word (.docx).

- template="cv5512" (mặc định): build bằng python-docx theo đúng khung Phụ lục IV,
  tái dùng `add_math_to_paragraph` để render công thức Toán ($...$) thành phương
  trình Word (OMML). KHÔNG cần docxtpl.
- template="school": điền template .docx của trường bằng docxtpl (M5, cần upload
  template + cài docxtpl). Hiện trả lỗi rõ ràng nếu chưa cấu hình.

LLM chỉ sinh JSON; toàn bộ format Word nằm ở đây (template-driven, chống lỗi format).
"""

from __future__ import annotations

import io
from typing import Any, Dict, List, Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.services.exporter import AI_DISCLOSURE_TEXT
from app.services.latex_to_omml import add_math_to_paragraph
from app.services.lesson_plan_schema import HOAT_DONG_LABELS


def render_khbd_docx(
    content: Dict[str, Any],
    template: str = "cv5512",
    template_path: Optional[str] = None,
) -> bytes:
    """Trả về bytes của file .docx."""
    if template == "school":
        return _render_school(content, template_path)
    return _render_cv5512(content)


# ── Default CV5512 (python-docx + OMML) ─────────────────────────────────────

def _p(doc, text: str = "", *, bold=False, size=None, align=None, italic=False):
    """Paragraph thường (không math)."""
    para = doc.add_paragraph()
    if align is not None:
        para.alignment = align
    if text:
        run = para.add_run(text)
        run.bold = bold
        run.italic = italic
        if size:
            run.font.size = Pt(size)
    return para


def _math_p(doc, label: str = "", value: str = "", *, indent=0, bullet=False):
    """Paragraph có thể chứa công thức Toán trong `value` ($...$)."""
    para = doc.add_paragraph()
    if indent:
        para.paragraph_format.left_indent = Pt(indent)
    prefix = ""
    if bullet:
        prefix = "• "
    if label:
        run = para.add_run(prefix + label)
        run.bold = True
    elif prefix:
        para.add_run(prefix)
    if value:
        add_math_to_paragraph(para, value)
    return para


def _heading(doc, text: str):
    return _p(doc, text, bold=True, size=13)


def _render_cv5512(content: Dict[str, Any]) -> bytes:
    doc = Document()
    meta = content.get("meta", {})
    mt = content.get("muc_tieu", {})
    tt = content.get("tien_trinh", []) or []

    # ── Tiêu đề ──
    ten_bai = meta.get("ten_bai", "KẾ HOẠCH BÀI DẠY")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run(f"TÊN BÀI DẠY: {ten_bai}".upper())
    tr.bold = True
    tr.font.size = Pt(14)

    mon = meta.get("mon", "Toán")
    lop = meta.get("lop", "")
    bo_sach = meta.get("bo_sach", "Kết nối tri thức")
    _p(doc, f"Môn học/Hoạt động giáo dục: {mon}; Lớp: {lop} — Bộ sách {bo_sach}")
    _p(doc, f"Thời gian thực hiện: {meta.get('so_tiet', 1)} tiết", italic=True)

    # ── I. MỤC TIÊU ──
    _heading(doc, "I. MỤC TIÊU")

    _p(doc, "1. Về kiến thức", bold=True)
    for kt in (mt.get("kien_thuc") or []):
        _math_p(doc, value=kt, bullet=True, indent=18)

    _p(doc, "2. Về năng lực", bold=True)
    nl = mt.get("nang_luc", {}) or {}
    _p(doc, "a) Năng lực chung", bold=True)
    for item in (nl.get("chung") or []):
        _nang_luc_line(doc, item)
    _p(doc, "b) Năng lực đặc thù", bold=True)
    for item in (nl.get("dac_thu") or []):
        _nang_luc_line(doc, item)

    _p(doc, "3. Về phẩm chất", bold=True)
    for item in (mt.get("pham_chat") or []):
        _nang_luc_line(doc, item)

    # ── II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU ──
    _heading(doc, "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU")
    for tb in (content.get("thiet_bi_day_hoc") or []):
        _math_p(doc, value=tb, bullet=True, indent=18)

    # ── III. TIẾN TRÌNH DẠY HỌC ──
    _heading(doc, "III. TIẾN TRÌNH DẠY HỌC")
    for i, hd in enumerate(tt, 1):
        _render_hoat_dong(doc, i, hd)

    # Dấu hiệu nhận dạng nội dung AI (Điều 44 Luật CN CNS + Luật AI). KHBD luôn
    # do AI sinh nên nhãn được gắn không điều kiện.
    _add_ai_disclosure(doc, ten_bai)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_ai_disclosure(doc, title: str) -> None:
    """Thêm dòng nhãn AI ở cuối KHBD + ghi docProps."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(AI_DISCLOSURE_TEXT)
    run.italic = True
    run.font.size = Pt(7.5)
    try:
        props = doc.core_properties
        props.title = title
        props.author = "Edu Smart App"
        props.comments = AI_DISCLOSURE_TEXT
        props.keywords = "AIGenerated=true"
        props.category = "AI-assisted"
    except Exception:  # pragma: no cover - metadata không được làm hỏng export
        pass


def _nang_luc_line(doc, item: Dict[str, Any]):
    ten = (item or {}).get("ten", "")
    bh = (item or {}).get("bieu_hien", "")
    _math_p(doc, label=f"{ten}: " if ten else "", value=bh, bullet=True, indent=18)


def _render_hoat_dong(doc, idx: int, hd: Dict[str, Any]):
    loai = hd.get("loai", "")
    label = HOAT_DONG_LABELS.get(loai, f"Hoạt động {idx}")
    ten = hd.get("ten_hoat_dong", "")
    tg = hd.get("thoi_gian_phut")
    head = label
    if ten and ten not in label:
        head += f" — {ten}"
    if tg:
        head += f" (khoảng {tg} phút)"
    _p(doc, head, bold=True, size=12)

    _math_p(doc, label="a) Mục tiêu: ", value=hd.get("muc_tieu", ""), indent=18)

    nd = hd.get("noi_dung", {}) or {}
    _math_p(doc, label="b) Nội dung: ", value=nd.get("mo_ta", ""), indent=18)
    for q in (nd.get("cau_hoi_nhiem_vu") or []):
        _math_p(doc, value=q, bullet=True, indent=36)

    sp = hd.get("san_pham", {}) or {}
    _math_p(doc, label="c) Sản phẩm: ", value=sp.get("mo_ta", ""), indent=18)
    for r in (sp.get("ket_qua_mong_doi") or []):
        _math_p(doc, value=r, bullet=True, indent=36)

    _p(doc, "d) Tổ chức thực hiện:", bold=True)
    toc = hd.get("to_chuc_thuc_hien", {}) or {}
    for key, lbl in [
        ("chuyen_giao_nhiem_vu", "Chuyển giao nhiệm vụ"),
        ("thuc_hien_nhiem_vu", "Thực hiện nhiệm vụ"),
        ("bao_cao_thao_luan", "Báo cáo, thảo luận"),
        ("ket_luan_nhan_dinh", "Kết luận, nhận định"),
    ]:
        val = (toc.get(key) or "").strip()
        if val:
            _math_p(doc, label=f"{lbl}: ", value=val, bullet=True, indent=36)


# ── School template (docxtpl) — M5 ──────────────────────────────────────────

def _render_school(content: Dict[str, Any], template_path: Optional[str]) -> bytes:
    if not template_path:
        raise RuntimeError(
            "Chưa có template trường (.docx). Hãy upload template trước khi xuất."
        )
    try:
        from docxtpl import DocxTemplate  # optional dependency (M5)
    except ImportError:
        raise ImportError("Cần cài 'docxtpl' để dùng template trường.")

    tpl = DocxTemplate(template_path)
    # Context phẳng cho Jinja; công thức giữ dạng $...$ (text). Render math nâng cao
    # cho template trường để dành về sau.
    tpl.render(_flatten_for_jinja(content))
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def _flatten_for_jinja(content: Dict[str, Any]) -> Dict[str, Any]:
    meta = content.get("meta", {})
    return {
        "meta": meta,
        "muc_tieu": content.get("muc_tieu", {}),
        "thiet_bi_day_hoc": content.get("thiet_bi_day_hoc", []),
        "tien_trinh": [
            {**hd, "label": HOAT_DONG_LABELS.get(hd.get("loai", ""), "")}
            for hd in (content.get("tien_trinh", []) or [])
        ],
    }

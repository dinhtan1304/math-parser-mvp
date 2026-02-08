"""
File Handler - Extract text from PDF and DOC/DOCX files
Optimized for Vietnamese math documents

Priority libraries:
1. pymupdf (fitz) - Best for PDF with complex layouts
2. pdfplumber - Good fallback for tables
3. python-docx - For DOCX files
"""

import os
import re
import base64
import asyncio
from pathlib import Path
from typing import Dict, Any, Optional, List
from concurrent.futures import ThreadPoolExecutor
import hashlib


class FileHandler:
    """
    Extract text from PDF and DOC/DOCX files.
    Optimized for Vietnamese math documents.
    """
    
    def __init__(self):
        self.executor = ThreadPoolExecutor(max_workers=4)
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check available libraries"""
        self.has_pymupdf = False
        self.has_pdfplumber = False
        self.has_pypdf = False
        self.has_docx = False
        
        try:
            import fitz
            self.has_pymupdf = True
            print("✅ pymupdf (fitz) available")
        except ImportError:
            print("⚠️ pymupdf not installed. Run: pip install pymupdf")
        
        try:
            import pdfplumber
            self.has_pdfplumber = True
        except ImportError:
            pass
        
        try:
            from pypdf import PdfReader
            self.has_pypdf = True
        except ImportError:
            pass
        
        try:
            from docx import Document
            self.has_docx = True
            print("✅ python-docx available")
        except ImportError:
            print("⚠️ python-docx not installed. Run: pip install python-docx")
    
    async def extract_text(self, file_path: str, use_vision: bool = False) -> Dict[str, Any]:
        """
        Extract text from file.
        
        Args:
            file_path: Path to file
            use_vision: If True, return images for Vision API
        
        Returns:
            {
                "text": str,
                "page_count": int,
                "file_type": str,
                "method": str,
                "images": List[Dict] (if use_vision)
            }
        """
        path = Path(file_path)
        ext = path.suffix.lower()
        
        file_hash = await self._compute_hash(file_path)
        
        if ext == '.pdf':
            if use_vision:
                result = await self._pdf_to_images(file_path)
            else:
                result = await self._extract_pdf(file_path)
        elif ext == '.docx':
            if use_vision:
                # Convert DOCX to PDF first, then to images
                result = await self._docx_to_images(file_path)
            else:
                result = await self._extract_docx(file_path)
        elif ext == '.doc':
            result = await self._extract_doc(file_path)
        elif ext in {'.txt', '.md'}:
            result = await self._extract_text_file(file_path)
        elif ext in {'.png', '.jpg', '.jpeg', '.gif', '.webp'}:
            result = await self._extract_image(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        result["file_hash"] = file_hash
        return result
    
    # ==================== PDF EXTRACTION ====================
    
    async def _extract_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF - tries multiple methods"""
        
        # Method 1: PyMuPDF (best)
        if self.has_pymupdf:
            result = await self._extract_pdf_pymupdf(file_path)
            if self._is_quality_good(result.get("text", "")):
                return result
            print("⚠️ PyMuPDF quality poor, trying pdfplumber...")
        
        # Method 2: pdfplumber
        if self.has_pdfplumber:
            result = await self._extract_pdf_pdfplumber(file_path)
            if self._is_quality_good(result.get("text", "")):
                return result
            print("⚠️ pdfplumber quality poor, trying pypdf...")
        
        # Method 3: pypdf (fallback)
        if self.has_pypdf:
            result = await self._extract_pdf_pypdf(file_path)
            return result
        
        return {"text": "", "error": "No PDF library available", "file_type": "pdf", "page_count": 0}
    
    async def _extract_pdf_pymupdf(self, file_path: str) -> Dict[str, Any]:
        """Extract using PyMuPDF - BEST for math documents"""
        loop = asyncio.get_event_loop()
        
        def extract():
            import fitz
            
            doc = fitz.open(file_path)
            text_parts = []
            
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # Extract with multiple methods and choose best
                # Method A: blocks (preserves layout)
                text_blocks = page.get_text("blocks")
                block_text = "\n".join([b[4] for b in text_blocks if b[6] == 0])
                
                # Method B: simple text
                simple_text = page.get_text("text")
                
                # Use the one with better quality
                if len(block_text) >= len(simple_text) * 0.9:
                    page_text = block_text
                else:
                    page_text = simple_text
                
                if page_text.strip():
                    text_parts.append(f"[Trang {page_num + 1}]")
                    text_parts.append(page_text)
            
            page_count = len(doc)
            doc.close()
            
            return "\n\n".join(text_parts), page_count
        
        text, page_count = await loop.run_in_executor(self.executor, extract)
        text = self._clean_text(text)
        
        return {
            "text": text,
            "page_count": page_count,
            "file_type": "pdf",
            "method": "pymupdf"
        }
    
    async def _extract_pdf_pdfplumber(self, file_path: str) -> Dict[str, Any]:
        """Extract using pdfplumber"""
        loop = asyncio.get_event_loop()
        
        def extract():
            import pdfplumber
            
            text_parts = []
            page_count = 0
            
            with pdfplumber.open(file_path) as pdf:
                page_count = len(pdf.pages)
                
                for i, page in enumerate(pdf.pages):
                    text = page.extract_text(
                        x_tolerance=2,
                        y_tolerance=2,
                        layout=True
                    )
                    if text:
                        text_parts.append(f"[Trang {i + 1}]")
                        text_parts.append(text)
            
            return "\n\n".join(text_parts), page_count
        
        text, page_count = await loop.run_in_executor(self.executor, extract)
        text = self._clean_text(text)
        
        return {
            "text": text,
            "page_count": page_count,
            "file_type": "pdf",
            "method": "pdfplumber"
        }
    
    async def _extract_pdf_pypdf(self, file_path: str) -> Dict[str, Any]:
        """Extract using pypdf (basic fallback)"""
        loop = asyncio.get_event_loop()
        
        def extract():
            from pypdf import PdfReader
            
            reader = PdfReader(file_path)
            text_parts = []
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Trang {i + 1}]")
                    text_parts.append(text)
            
            return "\n\n".join(text_parts), len(reader.pages)
        
        text, page_count = await loop.run_in_executor(self.executor, extract)
        text = self._clean_text(text)
        
        return {
            "text": text,
            "page_count": page_count,
            "file_type": "pdf",
            "method": "pypdf"
        }
    
    async def _pdf_to_images(self, file_path: str) -> Dict[str, Any]:
        """Convert PDF to images for Vision API using PyMuPDF (no poppler needed!)"""
        loop = asyncio.get_event_loop()
        
        def convert():
            # Method 1: PyMuPDF (fitz) - PREFERRED, no external dependencies
            if self.has_pymupdf:
                import fitz
                
                doc = fitz.open(file_path)
                base64_images = []
                
                # DPI ~150: zoom = 150/72 ≈ 2.08
                zoom = 2.0
                matrix = fitz.Matrix(zoom, zoom)
                
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    pix = page.get_pixmap(matrix=matrix)
                    img_bytes = pix.tobytes("jpeg")
                    b64 = base64.b64encode(img_bytes).decode()
                    
                    base64_images.append({
                        "page": page_num + 1,
                        "data": b64,
                        "mime_type": "image/jpeg"
                    })
                
                doc.close()
                print(f"✅ PyMuPDF rendered {len(base64_images)} pages to images")
                return base64_images, len(base64_images)
            
            # Method 2: pdf2image (needs poppler) - FALLBACK
            try:
                from pdf2image import convert_from_path
                import io
                
                images = convert_from_path(file_path, dpi=150, fmt='jpeg')
                base64_images = []
                
                for i, img in enumerate(images):
                    buffer = io.BytesIO()
                    img.save(buffer, format='JPEG', quality=85)
                    b64 = base64.b64encode(buffer.getvalue()).decode()
                    base64_images.append({
                        "page": i + 1,
                        "data": b64,
                        "mime_type": "image/jpeg"
                    })
                
                print(f"✅ pdf2image rendered {len(base64_images)} pages to images")
                return base64_images, len(base64_images)
                
            except ImportError:
                print("❌ Neither pymupdf nor pdf2image available!")
                return [], 0
        
        images, page_count = await loop.run_in_executor(self.executor, convert)
        
        return {
            "text": "",
            "images": images,
            "page_count": page_count,
            "file_type": "pdf",
            "method": "vision-pymupdf" if self.has_pymupdf else "vision-pdf2image"
        }
    
    # ==================== DOCX EXTRACTION ====================
    
    async def _docx_to_images(self, file_path: str) -> Dict[str, Any]:
        """Convert DOCX to images via intermediate PDF using LibreOffice"""
        loop = asyncio.get_event_loop()
        
        def convert():
            import subprocess
            import tempfile
            
            # Try converting DOCX → PDF → images using LibreOffice
            try:
                with tempfile.TemporaryDirectory() as tmpdir:
                    result = subprocess.run(
                        ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir', tmpdir, file_path],
                        capture_output=True, text=True, timeout=60
                    )
                    
                    if result.returncode == 0:
                        # Find the generated PDF
                        pdf_name = Path(file_path).stem + '.pdf'
                        pdf_path = os.path.join(tmpdir, pdf_name)
                        
                        if os.path.exists(pdf_path) and self.has_pymupdf:
                            import fitz
                            doc = fitz.open(pdf_path)
                            base64_images = []
                            zoom = 2.0
                            matrix = fitz.Matrix(zoom, zoom)
                            
                            for page_num in range(len(doc)):
                                page = doc[page_num]
                                pix = page.get_pixmap(matrix=matrix)
                                img_bytes = pix.tobytes("jpeg")
                                b64 = base64.b64encode(img_bytes).decode()
                                base64_images.append({
                                    "page": page_num + 1,
                                    "data": b64,
                                    "mime_type": "image/jpeg"
                                })
                            
                            doc.close()
                            print(f"✅ DOCX → PDF → {len(base64_images)} images")
                            return base64_images, len(base64_images)
            except Exception as e:
                print(f"⚠️ LibreOffice conversion failed: {e}")
            
            # Fallback: extract text instead
            print("⚠️ Cannot convert DOCX to images, falling back to text extraction")
            return [], 0
        
        images, page_count = await loop.run_in_executor(self.executor, convert)
        
        if images:
            return {
                "text": "",
                "images": images,
                "page_count": page_count,
                "file_type": "docx",
                "method": "vision-libreoffice"
            }
        else:
            # Fallback to text extraction
            return await self._extract_docx(file_path)
    
    async def _extract_docx(self, file_path: str) -> Dict[str, Any]:
        """Extract text from DOCX"""
        if not self.has_docx:
            return {"text": "", "error": "python-docx not installed", "file_type": "docx", "page_count": 0}
        
        loop = asyncio.get_event_loop()
        
        def extract():
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    text_parts.append(text)
            
            # Extract from tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        table_rows.append(" | ".join(cells))
                if table_rows:
                    text_parts.append("\n".join(table_rows))
            
            return "\n\n".join(text_parts)
        
        text = await loop.run_in_executor(self.executor, extract)
        text = self._clean_text(text)
        
        return {
            "text": text,
            "page_count": 1,
            "file_type": "docx",
            "method": "python-docx"
        }
    
    async def _extract_doc(self, file_path: str) -> Dict[str, Any]:
        """Extract text from old .doc file"""
        import subprocess
        
        loop = asyncio.get_event_loop()
        
        def extract():
            # Try antiword
            try:
                result = subprocess.run(
                    ['antiword', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout:
                    return result.stdout, "antiword"
            except:
                pass
            
            # Try catdoc
            try:
                result = subprocess.run(
                    ['catdoc', file_path],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0 and result.stdout:
                    return result.stdout, "catdoc"
            except:
                pass
            
            return "", "none"
        
        text, method = await loop.run_in_executor(self.executor, extract)
        
        if not text:
            return {
                "text": "",
                "error": "Cannot extract .doc file. Install antiword: apt install antiword",
                "file_type": "doc",
                "page_count": 0
            }
        
        text = self._clean_text(text)
        
        return {
            "text": text,
            "page_count": 1,
            "file_type": "doc",
            "method": method
        }
    
    # ==================== OTHER FORMATS ====================
    
    async def _extract_text_file(self, file_path: str) -> Dict[str, Any]:
        """Read plain text file"""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    text = f.read()
                return {
                    "text": text,
                    "page_count": 1,
                    "file_type": "text",
                    "method": encoding
                }
            except UnicodeDecodeError:
                continue
        
        return {"text": "", "error": "Could not decode text file", "file_type": "text", "page_count": 0}
    
    async def _extract_image(self, file_path: str) -> Dict[str, Any]:
        """Convert image to base64 for Vision API"""
        with open(file_path, 'rb') as f:
            b64 = base64.b64encode(f.read()).decode()
        
        ext = Path(file_path).suffix.lower()
        mime_map = {'.png': 'image/png', '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', 
                    '.gif': 'image/gif', '.webp': 'image/webp'}
        mime = mime_map.get(ext, 'image/jpeg')
        
        return {
            "text": "",
            "images": [{"page": 1, "data": b64, "mime_type": mime}],
            "page_count": 1,
            "file_type": "image",
            "method": "image"
        }
    
    # ==================== UTILITIES ====================
    
    def _is_quality_good(self, text: str) -> bool:
        """Check if extracted text quality is acceptable"""
        if not text or len(text) < 100:
            return False
        
        # Too many newlines = broken layout
        newline_ratio = text.count('\n') / len(text)
        if newline_ratio > 0.2:
            return False
        
        # Too many single-char lines
        lines = [l for l in text.split('\n') if l.strip()]
        single_char = sum(1 for l in lines if len(l.strip()) <= 2)
        if lines and single_char / len(lines) > 0.3:
            return False
        
        return True
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize text"""
        if not text:
            return ""
        
        # Remove control characters
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', text)
        
        # Normalize whitespace
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        text = re.sub(r' {3,}', '  ', text)
        text = re.sub(r'\t+', ' ', text)
        
        # Fix Vietnamese encoding issues
        text = text.replace('Ð', 'Đ').replace('ð', 'đ')
        
        return text.strip()
    
    async def _compute_hash(self, file_path: str) -> str:
        """Compute MD5 hash of file"""
        hash_md5 = hashlib.md5()
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(8192), b''):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()


# ==================== TEST ====================

if __name__ == "__main__":
    import sys
    
    async def test():
        handler = FileHandler()
        
        if len(sys.argv) > 1:
            result = await handler.extract_text(sys.argv[1])
            print(f"\nMethod: {result.get('method')}")
            print(f"Pages: {result.get('page_count')}")
            print(f"Text length: {len(result.get('text', ''))}")
            if result.get('error'):
                print(f"Error: {result['error']}")
            print(f"\n--- Preview ---\n{result.get('text', '')[:2000]}")
        else:
            print("Usage: python file_handler.py <file.pdf>")
    
    asyncio.run(test())
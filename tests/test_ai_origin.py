"""Gắn nhãn nguồn gốc nội dung AI (Điều 44 Luật CN CNS + Luật AI + NĐ 142/2026).

Nguyên tắc kiểm thử: KHÔNG bao giờ được suy đoán nhãn AI từ dữ liệu — nội dung do
người soạn phải giữ HUMAN, nội dung AI phải mang nhãn AI và nhãn đó phải theo bản
sao khi câu hỏi được clone sang ngân hàng người khác.
"""
import io

from app.core import content_origin
from app.services.exporter import AI_DISCLOSURE_TEXT, export_docx

API = "/api/v1"


# ── Hằng số & helper ─────────────────────────────────────────────────────────

def test_ai_origins_include_ocr_but_not_human():
    assert content_origin.is_ai_origin(content_origin.AI_GENERATED)
    assert content_origin.is_ai_origin(content_origin.AI_ASSISTED)
    # OCR cũng là xử lý bằng máy → vẫn phải gắn nhãn cho người dùng biết.
    assert content_origin.is_ai_origin(content_origin.OCR_IMPORT)
    assert not content_origin.is_ai_origin(content_origin.HUMAN)
    assert not content_origin.is_ai_origin(None)


# ── Nhãn khi lưu câu hỏi ─────────────────────────────────────────────────────

def test_manual_bulk_save_defaults_to_human(client, make_teacher):
    """Lưu thủ công KHÔNG được tự động bị dán nhãn AI."""
    _, headers = make_teacher()
    client.post(f"{API}/questions/bulk", json={
        "questions": [{"question_text": "Câu hỏi do giáo viên tự soạn", "answer": "A"}],
    }, headers=headers)

    items = client.get(f"{API}/questions", headers=headers).json()["items"]
    assert items[0]["origin"] == content_origin.HUMAN
    assert items[0]["reviewed_by_user"] is False


def test_bulk_save_honours_declared_ai_origin(client, make_teacher):
    """Caller khai báo nguồn AI → nhãn được ghi nhận đúng."""
    _, headers = make_teacher()
    r = client.post(f"{API}/questions/bulk", json={
        "questions": [{"question_text": "Câu hỏi do AI sinh ra", "answer": "B"}],
        "origin": content_origin.AI_GENERATED,
        "ai_model": "gemini-2.5-flash",
    }, headers=headers)
    assert r.status_code in (200, 201), r.text

    items = client.get(f"{API}/questions", headers=headers).json()["items"]
    saved = next(i for i in items if "AI sinh ra" in i["question_text"])
    assert saved["origin"] == content_origin.AI_GENERATED
    assert saved["ai_model"] == "gemini-2.5-flash"


def test_bulk_save_rejects_unknown_origin(client, make_teacher):
    _, headers = make_teacher()
    r = client.post(f"{API}/questions/bulk", json={
        "questions": [{"question_text": "X"}],
        "origin": "KHONG_TON_TAI",
    }, headers=headers)
    assert r.status_code == 400


def test_clone_preserves_ai_label(client, make_teacher):
    """Nhãn AI phải đi theo nội dung khi được sao chép sang ngân hàng khác."""
    _, owner_headers = make_teacher()
    client.post(f"{API}/questions/bulk", json={
        "questions": [{"question_text": "Câu AI chia sẻ cộng đồng", "answer": "C"}],
        "origin": content_origin.AI_GENERATED,
        "ai_model": "gemini-2.5-flash",
    }, headers=owner_headers)

    items = client.get(f"{API}/questions", headers=owner_headers).json()["items"]
    qid = items[0]["id"]
    client.patch(f"{API}/questions/bulk-visibility", json={
        "question_ids": [qid], "is_public": True,
    }, headers=owner_headers)

    _, other_headers = make_teacher()
    r = client.post(f"{API}/questions/{qid}/clone", headers=other_headers)
    assert r.status_code in (200, 201), r.text

    cloned = client.get(f"{API}/questions", headers=other_headers).json()["items"]
    assert cloned, "bản sao phải nằm trong ngân hàng người nhận"
    assert cloned[0]["origin"] == content_origin.AI_GENERATED
    assert cloned[0]["ai_model"] == "gemini-2.5-flash"


# ── Nhãn trong file xuất ra ──────────────────────────────────────────────────

def _docx_text(buf: io.BytesIO) -> str:
    from docx import Document
    buf.seek(0)
    return "\n".join(p.text for p in Document(buf).paragraphs)


def test_export_docx_labels_ai_content():
    buf = export_docx(
        [{"question": "Câu do AI sinh", "answer": "A", "origin": content_origin.AI_GENERATED}],
        title="ĐỀ KIỂM TRA",
    )
    assert AI_DISCLOSURE_TEXT in _docx_text(buf)


def test_export_docx_omits_label_for_human_content():
    """Nội dung do người soạn KHÔNG bị dán nhãn AI thừa."""
    buf = export_docx(
        [{"question": "Câu giáo viên tự soạn", "answer": "A", "origin": content_origin.HUMAN}],
        title="ĐỀ KIỂM TRA",
    )
    assert AI_DISCLOSURE_TEXT not in _docx_text(buf)


def test_export_docx_sets_ai_metadata():
    """File Word mang dấu hiệu nhận dạng cả trong metadata, không chỉ phần nhìn thấy."""
    from docx import Document

    buf = export_docx(
        [{"question": "Câu do AI sinh", "answer": "A", "origin": content_origin.AI_GENERATED}],
        title="ĐỀ KIỂM TRA",
    )
    buf.seek(0)
    props = Document(buf).core_properties
    assert props.keywords == "AIGenerated=true"
    assert AI_DISCLOSURE_TEXT in (props.comments or "")

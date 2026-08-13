"""
Test đặc tả cho `app/services/latex_to_omml.py` — bộ chuyển LaTeX → OMML.

TẠI SAO FILE NÀY QUAN TRỌNG
Module này quyết định công thức toán hiển thị thế nào trong file Word xuất ra,
tức là thứ người dùng cuối nhìn thấy. Trước 2026-08-12 nó KHÔNG có test nào
(`tests/test_latex_utils.py` test module KHÁC: `app/benchmark/latex_utils.py`).

ĐÂY LÀ TEST ĐẶC TẢ (characterization test): mục tiêu là KHÓA hành vi hiện tại
để refactor không làm đổi output, KHÔNG phải mô tả hành vi lý tưởng. Chỗ nào
hành vi hiện tại sai, dùng `xfail(strict=True)` với assertion ĐÚNG — khi ai đó
sửa lỗi, test sẽ XPASS và báo đỏ, buộc gỡ marker.

Module này không phụ thuộc gì trong `app/` (chỉ re, lxml, docx) nên test chạy
thuần, không cần DB/fixture.
"""

import pytest
from docx import Document

from app.services.latex_to_omml import (
    LaTeXToOMML,
    _split_math,
    add_math_to_paragraph,
    latex_to_text,
)


# ─────────────────────────────────────────────────────────────
# _split_math — tách text thành đoạn toán / không toán
# ─────────────────────────────────────────────────────────────

def test_split_math_tach_dung_cong_thuc_inline():
    assert _split_math("Cho $x>0$ thi") == [
        (False, "Cho "),
        (True, "x>0"),
        (False, " thi"),
    ]


def test_split_math_text_thuan_tra_ve_mot_doan():
    assert _split_math("Khong co cong thuc") == [(False, "Khong co cong thuc")]


def test_split_math_bo_dau_dollar_kep():
    # $$...$$ (display math) cũng được bóc, trả về nội dung đã strip
    assert _split_math(r"$$\frac{a}{b}$$") == [(True, r"\frac{a}{b}")]


def test_split_math_nhieu_cong_thuc_xen_ke():
    assert _split_math("A $x$ B $y$ C") == [
        (False, "A "),
        (True, "x"),
        (False, " B "),
        (True, "y"),
        (False, " C"),
    ]


def test_split_math_chuoi_rong():
    assert _split_math("") == [(False, "")]


def test_split_math_cong_thuc_rong_van_tinh_la_toan():
    # "$ $" → đoạn toán rỗng. Giữ hành vi này vì _insert_omml có nhánh
    # fallback riêng cho trường hợp convert() trả về rỗng.
    assert _split_math("Gia $ $ rong") == [
        (False, "Gia "),
        (True, ""),
        (False, " rong"),
    ]


# ─────────────────────────────────────────────────────────────
# LaTeXToOMML.convert — sinh phần tử OMML
# ─────────────────────────────────────────────────────────────

def _tags(elements):
    """Tên thẻ OMML (bỏ namespace) của danh sách phần tử."""
    return [el.tag.split("}")[-1] for el in elements]


def test_convert_bieu_thuc_don_gian_thanh_cac_run():
    # "x+1" → ba run OMML (toán hạng / toán tử / toán hạng)
    assert _tags(LaTeXToOMML().convert("x+1")) == ["r", "r", "r"]


def test_convert_phan_so_thanh_the_f():
    assert _tags(LaTeXToOMML().convert(r"\frac{a}{b}")) == ["f"]


def test_convert_luy_thua_thanh_sSup():
    assert _tags(LaTeXToOMML().convert("x^2")) == ["sSup"]


def test_convert_chi_so_duoi_thanh_sSub():
    assert _tags(LaTeXToOMML().convert("x_1")) == ["sSub"]


def test_convert_can_thuc_thanh_rad():
    assert _tags(LaTeXToOMML().convert(r"\sqrt{2}")) == ["rad"]


def test_convert_phan_so_long_nhau_van_la_mot_the_f():
    # Cấu trúc lồng phải gói trong MỘT thẻ <m:f>, không bị bung phẳng
    assert _tags(LaTeXToOMML().convert(r"\frac{\sqrt{a}}{b}")) == ["f"]


def test_convert_ky_hieu_hy_lap_thanh_mot_run():
    assert _tags(LaTeXToOMML().convert(r"\alpha")) == ["r"]


def test_convert_chuoi_rong_tra_ve_rong():
    assert LaTeXToOMML().convert("") == []


# ─────────────────────────────────────────────────────────────
# add_math_to_paragraph — tích hợp với python-docx
# ─────────────────────────────────────────────────────────────

def test_add_math_sinh_dung_so_oMath_va_run():
    para = Document().add_paragraph()
    add_math_to_paragraph(para, r"Cho $x^2+1$ va $\frac{a}{b}$ het")

    xml = para._element.xml
    # Hai công thức → hai khối <m:oMath>; ba đoạn text thường → ba <w:r>
    assert xml.count("<m:oMath>") == 2
    assert xml.count("<w:r>") == 3
    # para.text chỉ gom text thường, KHÔNG gồm nội dung công thức
    assert para.text == "Cho  va  het"


def test_add_math_text_thuan_khong_sinh_oMath():
    para = Document().add_paragraph()
    add_math_to_paragraph(para, "Khong co cong thuc")

    assert "oMath" not in para._element.xml
    assert para.text == "Khong co cong thuc"


def test_add_math_chuoi_rong_khong_them_gi():
    para = Document().add_paragraph()
    add_math_to_paragraph(para, "")
    assert len(para._element) == 0


def test_add_math_latex_hong_khong_nem_loi():
    """LaTeX sai cú pháp KHÔNG được làm vỡ luồng xuất file.

    Đây là ràng buộc quan trọng nhất của module: OCR sinh ra LaTeX không phải
    lúc nào cũng đúng, và một công thức hỏng không được phép làm hỏng cả file
    Word của giáo viên.
    """
    para = Document().add_paragraph()
    add_math_to_paragraph(para, r"$\frac{a$")  # thiếu ngoặc đóng

    # Không nổ. Hiện tại vẫn sinh ra <m:oMath> (rỗng) chứ không rơi về text thô.
    assert "oMath" in para._element.xml


def test_add_math_ap_dung_dinh_dang_cho_text_thuong():
    para = Document().add_paragraph()
    add_math_to_paragraph(para, "Dam $x$ nghieng", bold=True, font_size=12)

    xml = para._element.xml
    assert "<w:b/>" in xml or "<w:b " in xml


# ─────────────────────────────────────────────────────────────
# latex_to_text — fallback Unicode (dùng khi không xuất OMML được)
# ─────────────────────────────────────────────────────────────

def test_latex_to_text_phan_so():
    assert latex_to_text(r"$\frac{a}{b}$") == "(a)/(b)"


def test_latex_to_text_luy_thua_va_chi_so():
    assert latex_to_text("$x^2 + y_1$") == "x² + y₁"


def test_latex_to_text_can_bac_hai():
    assert latex_to_text(r"$\sqrt{16}$") == "√(16)"


def test_latex_to_text_ky_hieu_hy_lap():
    assert latex_to_text(r"$\alpha + \beta$") == "α + β"


def test_latex_to_text_giu_nguyen_text_ngoai_cong_thuc():
    assert latex_to_text(r"Cho $x^2$ va $\frac{1}{2}$ nhe") == "Cho x² va (1)/(2) nhe"


def test_latex_to_text_khong_co_cong_thuc_thi_giu_nguyen():
    assert latex_to_text("khong co gi") == "khong co gi"


def test_latex_to_text_chuoi_rong():
    assert latex_to_text("") == ""


# ─────────────────────────────────────────────────────────────
# HỒI QUY — hai lỗi tìm ra khi viết bộ test này, đã sửa 2026-08-12.
# Trước đó cả hai được đánh dấu xfail(strict=True) với assertion ĐÚNG; sửa
# xong chúng XPASS và báo đỏ, buộc gỡ marker. Cơ chế đó hoạt động như thiết kế.
# ─────────────────────────────────────────────────────────────

def test_lenh_dai_khong_bi_lenh_ngan_cat_cut():
    """LATEX_SYMBOLS phải được duyệt theo độ dài GIẢM DẦN.

    Nếu duyệt theo thứ tự dict, lệnh ngắn ăn trước và cắt cụt lệnh dài trùng
    tiền tố, để lại ký tự rác ngay giữa công thức của giáo viên.
    Có 6 cặp trùng tiền tố trong bảng ký hiệu — test đủ cả 6.
    """
    assert latex_to_text(r"$\neq$") == "≠"          # từng ra "≠q"
    assert latex_to_text(r"$\leq$") == "≤"          # từng ra "≤q"
    assert latex_to_text(r"$\geq$") == "≥"          # từng ra "≥q"
    assert latex_to_text(r"$\leftarrow$") == "←"    # từng ra "≤ftarrow"
    assert latex_to_text(r"$\cdots$") == "⋯"        # từng ra "·s"
    assert latex_to_text(r"$\infty$") == "∞"        # từng ra "∈fty"


def test_lenh_ngan_van_dung_sau_khi_doi_thu_tu():
    """Sửa thứ tự duyệt không được làm hỏng chính các lệnh ngắn."""
    assert latex_to_text(r"$\le$") == "≤"
    assert latex_to_text(r"$\ge$") == "≥"
    assert latex_to_text(r"$\ne$") == "≠"
    assert latex_to_text(r"$\in$") == "∈"
    assert latex_to_text(r"$\cdot$") == "·"   # U+00B7, khác \cdots → ⋯ (U+22EF)


def test_hon_hop_lenh_dai_va_ngan_trong_mot_cong_thuc():
    assert latex_to_text(r"$\le \ge \neq$") == "≤ ≥ ≠"


def test_can_bac_n_giu_dung_bac():
    """\\sqrt[n]{x} phải giữ bậc n, không phải luôn ra 'ⁿ√'."""
    assert latex_to_text(r"$\sqrt[3]{8}$") == "³√(8)"
    assert latex_to_text(r"$\sqrt[4]{16}$") == "⁴√(16)"


def test_can_bac_hai_khong_bi_anh_huong():
    assert latex_to_text(r"$\sqrt{16}$") == "√(16)"


def test_khong_con_lenh_nao_bi_tien_to_nuot():
    """Canh gác: thêm ký hiệu mới trùng tiền tố cũng phải an toàn.

    Duyệt mọi cặp (ngắn, dài) trùng tiền tố trong bảng và khẳng định lệnh dài
    dịch ra ĐÚNG ký hiệu của nó, không sót đuôi.
    """
    from app.services.latex_to_omml import LATEX_SYMBOLS

    co_ky_hieu = {k: v for k, v in LATEX_SYMBOLS.items() if v}
    for ngan in co_ky_hieu:
        for dai in co_ky_hieu:
            if dai != ngan and dai.startswith(ngan):
                assert latex_to_text(f"${dai}$") == co_ky_hieu[dai], (
                    f"{dai!r} bị {ngan!r} cắt cụt"
                )
